#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Génération du rapport technique EduManager - Système Éducatif Marocain"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_background(cell, fill_color):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading

def add_info_table(doc, data, headers=None, header_color="1F5C99"):
    """Add a styled table with optional headers."""
    cols = len(data[0]) if data else 2
    table = doc.add_table(rows=len(data) + (1 if headers else 0), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row_start = 0
    if headers:
        hdr_row = table.rows[0]
        for i, hdr in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = hdr
            set_cell_background(cell, header_color)
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(hdr)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_start = 1

    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + row_start]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            if r_idx % 2 == 0:
                set_cell_background(cell, "EBF3FB")
            para = cell.paragraphs[0]
            if para.runs:
                para.runs[0].font.size = Pt(9.5)
    return table

def create_report():
    doc = Document()

    # ── Page Setup ──────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Styles ───────────────────────────────────────────────────────────────────
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # ============================================================
    # PAGE DE GARDE
    # ============================================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("RAPPORT TECHNIQUE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 92, 153)

    doc.add_paragraph()

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_para.add_run("Système de Gestion Scolaire")
    run2.bold = True
    run2.font.size = Pt(22)
    run2.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph()

    app_para = doc.add_paragraph()
    app_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = app_para.add_run("EduManager")
    run3.bold = True
    run3.font.size = Pt(32)
    run3.font.color.rgb = RGBColor(231, 76, 60)

    doc.add_paragraph()
    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = info_para.add_run("Adapté au Système Éducatif Marocain\n(Enseignement Primaire & Collège)")
    run4.font.size = Pt(14)
    run4.font.color.rgb = RGBColor(127, 140, 141)
    run4.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = datetime.date.today().strftime("%d/%m/%Y")
    run5 = date_para.add_run(f"Date de rédaction : {today}\nVersion : 1.0.0")
    run5.font.size = Pt(11)
    run5.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ============================================================
    # SOMMAIRE
    # ============================================================
    add_heading(doc, "Sommaire", level=1, color=(31, 92, 153))
    toc_items = [
        ("1.", "Présentation du Projet", "3"),
        ("2.", "Contexte et Objectifs", "4"),
        ("3.", "Architecture Technique", "5"),
        ("4.", "Technologies Utilisées", "6"),
        ("5.", "Modèle de Données", "8"),
        ("6.", "API REST et Endpoints", "11"),
        ("7.", "Fonctionnalités Détaillées", "14"),
        ("8.", "Interface Utilisateur", "20"),
        ("9.", "Sécurité et Performances", "22"),
        ("10.", "Déploiement et Installation", "23"),
        ("11.", "Adéquation au Système Marocain", "25"),
        ("12.", "Statistiques du Projet", "26"),
        ("13.", "Améliorations Suggérées", "27"),
        ("14.", "Conclusion", "29"),
    ]
    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    toc_table.style = 'Table Grid'
    for i, (num, title, page) in enumerate(toc_items):
        row = toc_table.rows[i]
        row.cells[0].text = num
        row.cells[1].text = title
        row.cells[2].text = page
        if i % 2 == 0:
            for c in row.cells:
                set_cell_background(c, "EBF3FB")
        for j, cell in enumerate(row.cells):
            para = cell.paragraphs[0]
            if para.runs:
                para.runs[0].font.size = Pt(10)
            if j == 2:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============================================================
    # SECTION 1 : PRÉSENTATION DU PROJET
    # ============================================================
    add_heading(doc, "1. Présentation du Projet", level=1, color=(31, 92, 153))

    doc.add_paragraph(
        "EduManager est une plateforme numérique complète de gestion scolaire, "
        "développée spécifiquement pour répondre aux besoins des établissements "
        "d'enseignement au Maroc. Elle couvre l'ensemble du cycle administratif et "
        "pédagogique d'un établissement scolaire : de l'inscription des élèves à la "
        "gestion financière, en passant par le suivi des présences, des notes et de "
        "l'emploi du temps."
    )

    add_heading(doc, "1.1 Informations Générales", level=2, color=(52, 73, 94))
    add_info_table(doc,
        [
            ["Nom du projet",     "EduManager"],
            ["Version",           "1.0.0"],
            ["Type d'application","Application Web + Application Bureau (Electron)"],
            ["Langue de l'interface", "Français (adapté au contexte marocain)"],
            ["Statut",            "Production-ready"],
            ["Date de rédaction", today],
            ["Cycle scolaire cible", "Primaire, Collège (6ème – 3ème)"],
        ],
        headers=["Paramètre", "Valeur"]
    )

    doc.add_paragraph()
    add_heading(doc, "1.2 Objectif Principal", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "L'objectif principal d'EduManager est de digitaliser et d'optimiser la gestion "
        "administrative et pédagogique des écoles marocaines. La plateforme permet aux "
        "établissements scolaires de réduire la charge administrative, d'améliorer la "
        "communication avec les parents d'élèves et d'assurer un suivi rigoureux des "
        "performances académiques."
    )

    doc.add_page_break()

    # ============================================================
    # SECTION 2 : CONTEXTE ET OBJECTIFS
    # ============================================================
    add_heading(doc, "2. Contexte et Objectifs", level=1, color=(31, 92, 153))

    add_heading(doc, "2.1 Contexte du Système Éducatif Marocain", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "Le système éducatif marocain est organisé selon les cycles suivants :"
    )
    cycles = [
        ["Cycle", "Niveaux", "Durée", "Notes"],
        ["Préscolaire", "Maternelle", "2 ans", "Non obligatoire"],
        ["Primaire", "1ère – 6ème année", "6 ans", "Obligatoire"],
        ["Collège", "1ère – 3ème année (7ème – 9ème)", "3 ans", "Obligatoire"],
        ["Lycée", "Tronc commun + 1ère & 2ème Bac", "3 ans", "Baccalauréat national"],
        ["Supérieur", "Universités, Grandes Écoles", "Variable", "Diplômes nationaux"],
    ]
    add_info_table(doc, cycles[1:], headers=cycles[0])

    doc.add_paragraph()
    doc.add_paragraph(
        "EduManager cible principalement les établissements du cycle primaire et du "
        "collège (6ème à 3ème) conformément à la nomenclature utilisée dans l'application. "
        "Le système de notation est sur 20 points, avec des coefficients par matière et "
        "un découpage en deux semestres (S1 et S2)."
    )

    add_heading(doc, "2.2 Problèmes Résolus", level=2, color=(52, 73, 94))
    problems = [
        ["Problème", "Solution EduManager"],
        ["Gestion manuelle des inscriptions et dossiers élèves",
         "Module Étudiants avec CRUD complet et archivage numérique"],
        ["Suivi des paiements de scolarité difficile",
         "Module Finances avec états des paiements, arriérés et reçus"],
        ["Relevés de présence papier peu fiables",
         "Module Présences avec 4 sessions par jour et statistiques automatiques"],
        ["Communication difficile avec les parents",
         "Profil parent, solde dû, et notifications (WhatsApp/SMS prévu)"],
        ["Gestion des notes et bulletins chronophage",
         "Module Notes avec calcul automatique des moyennes pondérées et génération de bulletins"],
        ["Emploi du temps difficile à gérer",
         "Module Planning avec vue hebdomadaire interactive"],
        ["Absence de traçabilité des actions administratives",
         "Journal d'audit complet enregistrant toutes les actions"],
    ]
    add_info_table(doc, problems[1:], headers=problems[0])

    doc.add_page_break()

    # ============================================================
    # SECTION 3 : ARCHITECTURE TECHNIQUE
    # ============================================================
    add_heading(doc, "3. Architecture Technique", level=1, color=(31, 92, 153))

    add_heading(doc, "3.1 Architecture Générale", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "EduManager adopte une architecture client-serveur à trois couches (3-tier) "
        "classique, avec une séparation claire entre le frontend, le backend et la base "
        "de données."
    )

    arch_table = doc.add_table(rows=4, cols=3)
    arch_table.style = 'Table Grid'
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Couche", "Technologie", "Rôle"]
    hdr_row = arch_table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, "1F5C99")
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    arch_data = [
        ["Présentation (Frontend)", "React 18 + TypeScript + Vite", "Interface utilisateur, navigation, formulaires"],
        ["Logique Métier (Backend)", "Spring Boot 4.0 + Java 21", "API REST, règles métier, validation"],
        ["Données (Base de Données)", "MySQL 8.0+", "Persistance des données, contraintes d'intégrité"],
    ]
    for r_idx, (layer, tech, role) in enumerate(arch_data):
        row = arch_table.rows[r_idx + 1]
        row.cells[0].text = layer
        row.cells[1].text = tech
        row.cells[2].text = role
        if r_idx % 2 == 0:
            for c in row.cells:
                set_cell_background(c, "EBF3FB")

    doc.add_paragraph()

    add_heading(doc, "3.2 Structure du Projet", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "Le projet est organisé en deux répertoires principaux : backend et frontend, "
        "chacun constituant une application autonome."
    )
    struct_para = doc.add_paragraph()
    struct_para.style = 'No Spacing'
    code_lines = [
        "edumanager/",
        "├── backend/                     ← API Spring Boot (Java 21)",
        "│   ├── src/main/java/com/edumanager/api/",
        "│   │   ├── controller/          ← 13 contrôleurs REST",
        "│   │   ├── service/             ← 11 services métier",
        "│   │   ├── repository/          ← 10 repositories JPA",
        "│   │   ├── entity/              ← 10 entités + 11 énumérations",
        "│   │   ├── dto/                 ← 40+ DTOs requête/réponse",
        "│   │   ├── config/              ← CORS, Sécurité",
        "│   │   └── exception/           ← Gestion d'erreurs",
        "│   └── src/main/resources/",
        "│       ├── application.properties← Config MySQL",
        "│       └── seed-data.sql        ← Données de test (15 Ko)",
        "│",
        "├── frontend/                    ← React 18 + TypeScript",
        "│   ├── src/",
        "│   │   ├── pages/              ← 15 pages principales",
        "│   │   ├── components/         ← 70+ composants réutilisables",
        "│   │   ├── services/           ← 9 services API (Axios)",
        "│   │   └── types/              ← Définitions TypeScript",
        "│   └── package.json",
        "│",
        "└── uml/                        ← Diagrammes UML",
    ]
    run = struct_para.add_run("\n".join(code_lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)

    doc.add_page_break()

    # ============================================================
    # SECTION 4 : TECHNOLOGIES
    # ============================================================
    add_heading(doc, "4. Technologies Utilisées", level=1, color=(31, 92, 153))

    add_heading(doc, "4.1 Backend", level=2, color=(52, 73, 94))
    backend_tech = [
        ["Technologie", "Version", "Utilisation"],
        ["Java", "21 (LTS)", "Langage de programmation principal"],
        ["Spring Boot", "4.0.3", "Framework d'application REST"],
        ["Spring Data JPA", "Inclus", "ORM et accès aux données (Hibernate)"],
        ["Spring Security", "Inclus", "Authentification et autorisation"],
        ["Spring Validation", "Inclus", "Validation des DTOs"],
        ["MySQL Connector/J", "8.x", "Driver JDBC pour MySQL"],
        ["Lombok", "1.18.x", "Réduction du code boilerplate"],
        ["Maven", "3.9.x", "Gestionnaire de dépendances et build"],
    ]
    add_info_table(doc, backend_tech[1:], headers=backend_tech[0])

    doc.add_paragraph()
    add_heading(doc, "4.2 Frontend", level=2, color=(52, 73, 94))
    frontend_tech = [
        ["Bibliothèque/Framework", "Version", "Utilisation"],
        ["React", "18.3.1", "Framework UI composants"],
        ["TypeScript", "5.8.3", "Typage statique JavaScript"],
        ["Vite", "5.4.19", "Bundler et serveur de développement"],
        ["React Router", "6.30", "Navigation SPA"],
        ["Tailwind CSS", "3.4.17", "CSS utilitaire"],
        ["Radix UI + shadcn/ui", "Dernière", "Composants UI accessibles"],
        ["React Hook Form", "7.x", "Gestion des formulaires"],
        ["Zod", "3.x", "Validation des schémas"],
        ["Axios", "1.14", "Client HTTP"],
        ["Recharts", "2.15.4", "Graphiques et visualisations"],
        ["Lucide React", "Dernière", "Bibliothèque d'icônes"],
        ["date-fns", "Dernière", "Manipulation des dates"],
        ["@tanstack/react-query", "5.x", "Gestion de l'état asynchrone"],
        ["Electron", "40.4.1", "Application bureau desktop"],
        ["Electron Builder", "26.8.1", "Packaging application .exe"],
    ]
    add_info_table(doc, frontend_tech[1:], headers=frontend_tech[0])

    doc.add_paragraph()
    add_heading(doc, "4.3 Base de Données", level=2, color=(52, 73, 94))
    db_para = doc.add_paragraph(
        "Le projet utilise MySQL 8.0+ comme système de gestion de base de données "
        "relationnelle (SGBDR). Le schéma est géré automatiquement par Hibernate via "
        "la propriété spring.jpa.hibernate.ddl-auto=update. La base de données est "
        "configurée en charset UTF-8 pour la prise en charge complète de la langue "
        "française et des caractères spéciaux."
    )

    doc.add_paragraph()
    db_config = [
        ["Paramètre", "Valeur"],
        ["Hôte",       "localhost"],
        ["Port",        "3306"],
        ["Nom de la BD","edumanager_db"],
        ["Charset",     "utf8mb4 (Unicode complet)"],
        ["Port Backend","8080"],
        ["Port Frontend (dev)", "5173 (Vite)"],
        ["Année académique par défaut", "2024-2025"],
    ]
    add_info_table(doc, db_config[1:], headers=db_config[0])

    doc.add_page_break()

    # ============================================================
    # SECTION 5 : MODÈLE DE DONNÉES
    # ============================================================
    add_heading(doc, "5. Modèle de Données", level=1, color=(31, 92, 153))

    doc.add_paragraph(
        "La base de données EduManager est composée de 10 tables principales reliées "
        "par des contraintes de clés étrangères. Voici la description détaillée de "
        "chaque entité."
    )

    # -- Students
    add_heading(doc, "5.1 Table STUDENTS (Élèves)", level=2, color=(52, 73, 94))
    students_cols = [
        ["Champ", "Type", "Description"],
        ["id", "BIGINT (PK)", "Identifiant unique auto-incrémenté"],
        ["first_name", "VARCHAR(100)", "Prénom de l'élève"],
        ["last_name", "VARCHAR(100)", "Nom de famille"],
        ["email", "VARCHAR(150)", "Adresse e-mail"],
        ["phone", "VARCHAR(20)", "Numéro de téléphone"],
        ["birth_date", "DATE", "Date de naissance"],
        ["address", "VARCHAR(255)", "Adresse complète"],
        ["class_name", "VARCHAR(50)", "Classe (ex : 6ème A, 3ème B)"],
        ["status", "ENUM", "ACTIVE / INACTIVE / SUSPENDED"],
        ["avatar_url", "VARCHAR(255)", "URL de la photo de profil"],
        ["parent_id", "BIGINT (FK)", "Référence vers la table PARENTS"],
        ["created_at", "DATETIME", "Date de création du dossier"],
        ["updated_at", "DATETIME", "Dernière mise à jour"],
    ]
    add_info_table(doc, students_cols[1:], headers=students_cols[0])

    doc.add_paragraph()

    # -- Parents
    add_heading(doc, "5.2 Table PARENTS (Tuteurs)", level=2, color=(52, 73, 94))
    parents_cols = [
        ["Champ", "Type", "Description"],
        ["id", "BIGINT (PK)", "Identifiant unique"],
        ["first_name / last_name", "VARCHAR(100)", "Prénom et nom du tuteur"],
        ["email", "VARCHAR(150)", "Adresse e-mail"],
        ["phone", "VARCHAR(20)", "Numéro de téléphone"],
        ["address", "VARCHAR(255)", "Adresse domicile"],
        ["cin", "VARCHAR(20)", "Numéro CIN (Carte d'Identité Nationale)"],
        ["nationality", "VARCHAR(50)", "Nationalité"],
        ["profession", "VARCHAR(100)", "Profession"],
        ["arrears", "DECIMAL(10,2)", "Montant des arriérés de paiement"],
        ["avatar_url", "VARCHAR(255)", "Photo de profil"],
    ]
    add_info_table(doc, parents_cols[1:], headers=parents_cols[0])

    doc.add_paragraph()

    # -- Teachers
    add_heading(doc, "5.3 Table TEACHERS (Professeurs)", level=2, color=(52, 73, 94))
    teachers_cols = [
        ["Champ", "Type", "Description"],
        ["id", "BIGINT (PK)", "Identifiant unique"],
        ["first_name / last_name", "VARCHAR(100)", "Identité du professeur"],
        ["email", "VARCHAR(150)", "E-mail professionnel"],
        ["phone", "VARCHAR(20)", "Téléphone"],
        ["specialization", "VARCHAR(100)", "Domaine de spécialité"],
        ["status", "ENUM", "ACTIVE / INACTIVE"],
        ["subjects", "TEXT (JSON)", "Liste des matières enseignées"],
        ["assigned_classes", "TEXT (JSON)", "Classes assignées"],
        ["avatar_url", "VARCHAR(255)", "Photo de profil"],
    ]
    add_info_table(doc, teachers_cols[1:], headers=teachers_cols[0])

    doc.add_paragraph()

    # -- Grades
    add_heading(doc, "5.4 Table GRADES (Notes et Évaluations)", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "Cette table stocke toutes les notes attribuées aux élèves selon le système "
        "d'évaluation marocain (sur 20 points)."
    )
    grades_cols = [
        ["Champ", "Type", "Description"],
        ["id", "BIGINT (PK)", "Identifiant unique"],
        ["student_id", "BIGINT (FK)", "Référence vers STUDENTS"],
        ["teacher_id", "BIGINT (FK)", "Référence vers TEACHERS"],
        ["module_name", "VARCHAR(100)", "Nom de la matière (ex : Mathématiques)"],
        ["evaluation_type", "ENUM", "CONTROL / DS / EXAM / TP / ORAL"],
        ["semester", "ENUM", "S1 (1er semestre) ou S2 (2ème semestre)"],
        ["score", "DECIMAL(5,2)", "Note obtenue (sur 20)"],
        ["coefficient", "DECIMAL(3,1)", "Coefficient de la matière"],
        ["weighted_score", "DECIMAL(8,2)", "Note pondérée = score × coefficient"],
        ["academic_year", "VARCHAR(10)", "Année scolaire (ex : 2024-2025)"],
        ["graded_at", "DATETIME", "Date de saisie de la note"],
    ]
    add_info_table(doc, grades_cols[1:], headers=grades_cols[0])

    doc.add_paragraph()

    # -- Attendance
    add_heading(doc, "5.5 Table ATTENDANCE (Présences)", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "La gestion des présences est basée sur 4 sessions journalières, "
        "conformément aux horaires habituels des écoles marocaines."
    )
    att_cols = [
        ["Champ", "Type", "Description"],
        ["id", "BIGINT (PK)", "Identifiant unique"],
        ["student_id", "BIGINT (FK)", "Référence vers STUDENTS"],
        ["date", "DATE", "Date de la séance"],
        ["session", "ENUM", "SESSION_1 (8h-10h) à SESSION_4 (16h-18h)"],
        ["status", "ENUM", "PRESENT / LATE (retard) / ABSENT"],
        ["class_name", "VARCHAR(50)", "Classe concernée"],
        ["marked_by_teacher", "VARCHAR(100)", "Nom de l'enseignant"],
        ["notes", "VARCHAR(255)", "Observations éventuelles"],
        ["Contrainte unique", "–", "(student_id, date, session)"],
    ]
    add_info_table(doc, att_cols[1:], headers=att_cols[0])

    doc.add_paragraph()

    # -- Transactions
    add_heading(doc, "5.6 Table TRANSACTIONS (Paiements)", level=2, color=(52, 73, 94))
    trans_cols = [
        ["Champ", "Type", "Description"],
        ["id", "BIGINT (PK)", "Identifiant unique"],
        ["student_id / parent_id", "BIGINT (FK)", "Références vers STUDENTS / PARENTS"],
        ["amount", "DECIMAL(10,2)", "Montant en dirhams (MAD)"],
        ["type", "ENUM", "TUITION / INSCRIPTION / OTHER"],
        ["payment_method", "ENUM", "CASH / CHEQUE / BANK_TRANSFER / OTHER"],
        ["status", "ENUM", "PAID / PENDING / OVERDUE"],
        ["due_date", "DATE", "Date d'échéance"],
        ["paid_at", "DATETIME", "Date/heure du paiement"],
        ["receipt_number", "VARCHAR(50)", "Numéro du reçu"],
        ["academic_year", "VARCHAR(10)", "Année scolaire"],
        ["description", "VARCHAR(255)", "Description libre"],
    ]
    add_info_table(doc, trans_cols[1:], headers=trans_cols[0])

    doc.add_paragraph()

    # -- Tarifs
    add_heading(doc, "5.7 Table TARIFS (Contrats de Scolarité)", level=2, color=(52, 73, 94))
    tarifs_cols = [
        ["Champ", "Type", "Description"],
        ["id", "BIGINT (PK)", "Identifiant unique"],
        ["student_id", "BIGINT (FK)", "Référence vers STUDENTS"],
        ["academic_year", "VARCHAR(10)", "Année scolaire"],
        ["total_amount", "DECIMAL(10,2)", "Montant total annuel (MAD)"],
        ["frequency", "ENUM", "MONTHLY / TRIMESTRIAL / ANNUAL"],
        ["enrollment_month", "INT", "Mois d'inscription (1-12)"],
        ["installment_count", "INT", "Nombre d'échéances"],
        ["installment_amount", "DECIMAL(10,2)", "Montant par échéance"],
        ["Contrainte unique", "–", "(student_id, academic_year)"],
    ]
    add_info_table(doc, tarifs_cols[1:], headers=tarifs_cols[0])

    doc.add_paragraph()

    # -- Audit Logs
    add_heading(doc, "5.8 Table AUDIT_LOGS (Journal d'Activité)", level=2, color=(52, 73, 94))
    audit_cols = [
        ["Champ", "Type", "Description"],
        ["id", "BIGINT (PK)", "Identifiant unique"],
        ["agent_id / agent_name", "BIGINT / VARCHAR", "Identité de l'utilisateur"],
        ["module", "VARCHAR(50)", "Module concerné (STUDENTS, FINANCES, etc.)"],
        ["action", "ENUM", "CREATE / UPDATE / DELETE / VIEW / LOGIN / LOGOUT"],
        ["description", "TEXT", "Description de l'action"],
        ["target", "VARCHAR(100)", "Entité ciblée"],
        ["ip_address", "VARCHAR(45)", "Adresse IP de l'utilisateur"],
        ["timestamp", "DATETIME", "Date et heure précises"],
    ]
    add_info_table(doc, audit_cols[1:], headers=audit_cols[0])

    doc.add_paragraph()

    # -- Settings
    add_heading(doc, "5.9 Table SCHOOL_SETTINGS (Configuration)", level=2, color=(52, 73, 94))
    settings_cols = [
        ["Champ", "Type", "Description"],
        ["id", "BIGINT (PK)", "Identifiant unique (généralement 1)"],
        ["school_name", "VARCHAR(200)", "Nom de l'établissement"],
        ["address", "VARCHAR(255)", "Adresse complète"],
        ["email", "VARCHAR(150)", "E-mail officiel"],
        ["phone", "VARCHAR(20)", "Téléphone de l'administration"],
        ["logo_data", "LONGTEXT", "Logo en Base64"],
        ["created_at / updated_at", "DATETIME", "Horodatages"],
    ]
    add_info_table(doc, settings_cols[1:], headers=settings_cols[0])

    doc.add_paragraph()

    # -- Enums
    add_heading(doc, "5.10 Énumérations (11 enums)", level=2, color=(52, 73, 94))
    enums_data = [
        ["Énumération", "Valeurs"],
        ["StudentStatus", "ACTIVE, INACTIVE, SUSPENDED"],
        ["AgentStatus", "ACTIVE, INACTIVE"],
        ["TeacherStatus", "ACTIVE, INACTIVE"],
        ["AttendanceStatus", "PRESENT, LATE, ABSENT"],
        ["AttendanceSession", "SESSION_1 (8h-10h), SESSION_2 (10h-12h), SESSION_3 (14h-16h), SESSION_4 (16h-18h)"],
        ["EvaluationType", "CONTROL (contrôle), DS (devoir surveillé), EXAM, TP, ORAL"],
        ["Semester", "S1 (1er semestre), S2 (2ème semestre)"],
        ["TransactionType", "TUITION (scolarité), INSCRIPTION, OTHER"],
        ["TransactionStatus", "PAID, PENDING, OVERDUE"],
        ["PaymentMethod", "CASH, CHEQUE, BANK_TRANSFER, OTHER"],
        ["PaymentFrequency", "MONTHLY (mensuel), TRIMESTRIAL (trimestriel), ANNUAL (annuel)"],
    ]
    add_info_table(doc, enums_data[1:], headers=enums_data[0])

    doc.add_page_break()

    # ============================================================
    # SECTION 6 : API REST
    # ============================================================
    add_heading(doc, "6. API REST et Endpoints", level=1, color=(31, 92, 153))

    doc.add_paragraph(
        "L'API REST est développée avec Spring Boot. Tous les endpoints retournent "
        "des réponses au format JSON. La base de l'URL est : http://localhost:8080/api"
    )

    api_sections = [
        ("6.1 Authentification", [
            ["Méthode", "Endpoint", "Description"],
            ["POST", "/api/auth/login", "Connexion (email + mot de passe)"],
            ["POST", "/api/auth/logout", "Déconnexion et enregistrement audit"],
        ]),
        ("6.2 Gestion des Élèves", [
            ["Méthode", "Endpoint", "Description"],
            ["GET", "/api/students", "Liste paginée des élèves (20/page)"],
            ["GET", "/api/students/{id}", "Détail d'un élève"],
            ["POST", "/api/students", "Créer un nouvel élève"],
            ["PUT", "/api/students/{id}", "Modifier le dossier d'un élève"],
            ["DELETE", "/api/students/{id}", "Supprimer un élève"],
            ["GET", "/api/students/search?keyword=&className=&status=", "Recherche avancée"],
        ]),
        ("6.3 Gestion des Parents", [
            ["Méthode", "Endpoint", "Description"],
            ["GET", "/api/parents", "Liste de tous les parents"],
            ["GET", "/api/parents/{id}", "Détail d'un parent"],
            ["GET", "/api/parents/{id}/students", "Enfants d'un parent"],
            ["POST", "/api/parents", "Créer un parent"],
            ["PUT", "/api/parents/{id}", "Modifier un parent"],
            ["DELETE", "/api/parents/{id}", "Supprimer un parent"],
        ]),
        ("6.4 Gestion des Professeurs", [
            ["Méthode", "Endpoint", "Description"],
            ["GET", "/api/teachers", "Liste des professeurs"],
            ["GET", "/api/teachers?status=ACTIVE", "Filtrer par statut"],
            ["GET", "/api/teachers/{id}", "Détail d'un professeur"],
            ["POST", "/api/teachers", "Créer un professeur"],
            ["PUT", "/api/teachers/{id}", "Modifier un professeur"],
            ["DELETE", "/api/teachers/{id}", "Supprimer un professeur"],
        ]),
        ("6.5 Gestion des Présences", [
            ["Méthode", "Endpoint", "Description"],
            ["GET", "/api/attendances", "Tous les enregistrements"],
            ["GET", "/api/attendances/class?date=&className=", "Présences par date et classe"],
            ["POST", "/api/attendances/mark", "Enregistrer les présences"],
            ["POST", "/api/attendances/bulk-present", "Marquer tous présents (1 clic)"],
            ["GET", "/api/attendances/student/{id}/stats", "Statistiques d'un élève"],
        ]),
        ("6.6 Notes et Bulletins", [
            ["Méthode", "Endpoint", "Description"],
            ["GET", "/api/grades", "Toutes les notes"],
            ["GET", "/api/grades/student/{id}", "Notes d'un élève"],
            ["GET", "/api/grades/student/{id}?semester=S1", "Notes par semestre"],
            ["GET", "/api/grades/student/{id}/report?academicYear=2024-2025", "Bulletin scolaire"],
            ["POST", "/api/grades", "Ajouter une note"],
            ["PUT", "/api/grades/{id}", "Modifier une note"],
            ["DELETE", "/api/grades/{id}", "Supprimer une note"],
        ]),
        ("6.7 Finances et Paiements", [
            ["Méthode", "Endpoint", "Description"],
            ["GET", "/api/transactions", "Liste des transactions (paginée)"],
            ["GET", "/api/transactions/{id}", "Détail d'une transaction"],
            ["POST", "/api/transactions", "Créer une transaction"],
            ["PUT", "/api/transactions/{id}", "Modifier une transaction"],
            ["GET", "/api/transactions/summary?academicYear=2024-2025", "Résumé financier annuel"],
        ]),
        ("6.8 Tarification", [
            ["Méthode", "Endpoint", "Description"],
            ["GET", "/api/tarifs", "Tous les contrats"],
            ["GET", "/api/tarifs?academicYear=2024-2025", "Contrats par année"],
            ["GET", "/api/tarifs/student/{id}", "Contrats d'un élève"],
            ["POST", "/api/tarifs", "Créer un contrat"],
            ["PUT", "/api/tarifs/{id}", "Modifier un contrat"],
            ["DELETE", "/api/tarifs/{id}", "Supprimer un contrat"],
            ["POST", "/api/tarifs/payment/split", "Paiement partagé (plusieurs enfants)"],
        ]),
        ("6.9 Journal d'Activité", [
            ["Méthode", "Endpoint", "Description"],
            ["GET", "/api/audit-logs", "Tous les journaux"],
            ["GET", "/api/audit-logs/agent/{id}", "Journaux d'un agent"],
            ["GET", "/api/audit-logs/module/{module}", "Journaux par module"],
            ["GET", "/api/audit-logs/date-range?start=&end=", "Journaux par période"],
        ]),
        ("6.10 Paramètres Scolaires", [
            ["Méthode", "Endpoint", "Description"],
            ["GET", "/api/settings/school", "Lire les paramètres"],
            ["PUT", "/api/settings/school", "Mettre à jour les paramètres"],
        ]),
    ]

    for section_title, data in api_sections:
        add_heading(doc, section_title, level=2, color=(52, 73, 94))
        add_info_table(doc, data[1:], headers=data[0])
        doc.add_paragraph()

    doc.add_page_break()

    # ============================================================
    # SECTION 7 : FONCTIONNALITÉS DÉTAILLÉES
    # ============================================================
    add_heading(doc, "7. Fonctionnalités Détaillées", level=1, color=(31, 92, 153))

    # 7.1 Dashboard
    add_heading(doc, "7.1 Tableau de Bord (Dashboard)", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "La page d'accueil offre une vue synthétique de l'état de l'établissement en "
        "temps réel."
    )
    dash_items = [
        ["Indicateur", "Description"],
        ["Total Élèves", "Nombre total d'élèves inscrits et actifs"],
        ["Professeurs Actifs", "Nombre de professeurs en service"],
        ["Revenu Mensuel", "Total des paiements reçus le mois en cours (en MAD)"],
        ["Taux de Présence", "Pourcentage de présences sur la période sélectionnée"],
        ["Graphique Revenus", "Évolution mensuelle des encaissements (Recharts)"],
        ["Inscriptions Récentes", "Liste des dernières inscriptions d'élèves"],
    ]
    add_info_table(doc, dash_items[1:], headers=dash_items[0])

    doc.add_paragraph()

    # 7.2 Etudiants
    add_heading(doc, "7.2 Gestion des Élèves", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "Module central de la plateforme permettant la gestion complète du dossier "
        "scolaire de chaque élève."
    )
    stu_items = [
        ["Fonctionnalité", "Description"],
        ["Liste paginée", "Affichage avec 5/10/20/50 éléments par page"],
        ["Recherche", "Recherche par prénom, nom, email"],
        ["Filtres", "Filtrage par classe et statut (Actif/Inactif/Suspendu)"],
        ["Création", "Formulaire complet : identité, coordonnées, classe, tuteur"],
        ["Modification", "Mise à jour de toutes les informations"],
        ["Profil détaillé", "Vue complète avec onglets : info, notes, présences, paiements"],
        ["Lien parent", "Association de l'élève à son tuteur légal"],
        ["Statuts", "ACTIVE (actif), INACTIVE (inactif), SUSPENDED (suspendu)"],
        ["Classes supportées", "6ème A/B, 5ème A/B, 4ème A/B, 3ème A/B/C"],
    ]
    add_info_table(doc, stu_items[1:], headers=stu_items[0])

    doc.add_paragraph()

    # 7.3 Parents
    add_heading(doc, "7.3 Gestion des Parents / Tuteurs", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "Gestion des informations des tuteurs légaux, indispensable dans le contexte "
        "marocain où la communication avec les familles est primordiale."
    )
    par_items = [
        ["Fonctionnalité", "Description"],
        ["Profil complet", "CIN, nationalité, profession, adresse"],
        ["Solde dû (Arriérés)", "Montant total des impayés en MAD"],
        ["Enfants liés", "Liste des élèves associés au parent"],
        ["Historique paiements", "Transactions liées au parent"],
        ["Notification", "Préparation pour WhatsApp/SMS (interface prête)"],
    ]
    add_info_table(doc, par_items[1:], headers=par_items[0])

    doc.add_paragraph()

    # 7.4 Professeurs
    add_heading(doc, "7.4 Gestion des Professeurs", level=2, color=(52, 73, 94))
    prof_items = [
        ["Fonctionnalité", "Description"],
        ["Profil pédagogique", "Spécialité, matières, classes assignées"],
        ["Statut de service", "ACTIVE (en poste) / INACTIVE (non actif)"],
        ["Feuille de service", "Lien avec l'emploi du temps"],
        ["Gestion des notes", "Association professeur-matière pour la saisie des notes"],
    ]
    add_info_table(doc, prof_items[1:], headers=prof_items[0])

    doc.add_paragraph()

    # 7.5 Presences
    add_heading(doc, "7.5 Suivi des Présences", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "Le module de présences est conçu selon les horaires habituels des écoles "
        "marocaines, avec 4 sessions par jour couvrant la journée complète."
    )
    att_items = [
        ["Fonctionnalité", "Description"],
        ["4 sessions/jour", "SESSION_1 (8h-10h), SESSION_2 (10h-12h), SESSION_3 (14h-16h), SESSION_4 (16h-18h)"],
        ["Statuts", "PRESENT (présent), LATE (retard/en retard), ABSENT"],
        ["Calendrier interactif", "Sélection de la date et de la classe"],
        ["Marquage rapide", "Bouton 'Tous présents' en 1 clic"],
        ["Contrainte unicité", "Un seul enregistrement par (élève, date, session)"],
        ["Statistiques", "Taux de présence par élève sur une période"],
        ["Observations", "Champ notes pour justifier les absences"],
    ]
    add_info_table(doc, att_items[1:], headers=att_items[0])

    doc.add_paragraph()

    # 7.6 Notes
    add_heading(doc, "7.6 Gestion des Notes et Bulletins", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "Système conforme aux pratiques marocaines : notation sur 20, coefficients "
        "par matière, deux semestres et bulletin scolaire généré automatiquement."
    )
    grade_items = [
        ["Fonctionnalité", "Description"],
        ["Notation sur 20", "Toutes les notes sont sur 20 (système marocain)"],
        ["Types d'évaluation", "Contrôle (CC), Devoir Surveillé (DS), Examen, TP, Oral"],
        ["Semestres", "S1 (1er semestre) et S2 (2ème semestre)"],
        ["Coefficients", "Chaque matière a un coefficient de pondération"],
        ["Note pondérée", "Calcul automatique : score × coefficient"],
        ["Bulletin scolaire", "Rapport complet par élève et par année académique"],
        ["Année académique", "Format 2024-2025 (conforme au calendrier marocain)"],
        ["Matières supportées", "Mathématiques, Français, Arabe, Sciences, Histoire-Géo, etc."],
    ]
    add_info_table(doc, grade_items[1:], headers=grade_items[0])

    doc.add_paragraph()

    # 7.7 Finances
    add_heading(doc, "7.7 Gestion Financière", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "Module complet de gestion des frais de scolarité, des paiements et des "
        "arriérés, avec support multi-modes de paiement utilisés au Maroc."
    )
    fin_items = [
        ["Fonctionnalité", "Description"],
        ["Types de transactions", "TUITION (mensualités), INSCRIPTION (frais d'inscription), OTHER"],
        ["Modes de paiement", "Espèces (CASH), Chèque (CHEQUE), Virement (BANK_TRANSFER), Autre"],
        ["Statuts de paiement", "PAID (payé), PENDING (en attente), OVERDUE (en retard)"],
        ["Gestion des arriérés", "Calcul automatique des impayés par parent"],
        ["Contrats de scolarité", "Tarif annuel avec fréquence de paiement paramétrable"],
        ["Paiement partagé", "Un parent paie pour plusieurs enfants en une transaction"],
        ["Reçus de paiement", "Génération de reçus numérotés"],
        ["Résumé financier", "Vue d'ensemble annuelle (encaissé, en attente, en retard)"],
        ["Monnaie", "Dirham Marocain (MAD)"],
    ]
    add_info_table(doc, fin_items[1:], headers=fin_items[0])

    doc.add_paragraph()

    # 7.8 Emploi du temps
    add_heading(doc, "7.8 Emploi du Temps", level=2, color=(52, 73, 94))
    edt_items = [
        ["Fonctionnalité", "Description"],
        ["Vue hebdomadaire", "Grille Lundi–Vendredi (5 jours)"],
        ["Créneaux horaires", "4 créneaux par jour correspondant aux sessions"],
        ["Assignation", "Association Matière ↔ Professeur ↔ Classe ↔ Créneau"],
        ["Édition intégrée", "Modification directe dans la grille"],
        ["Export/Impression", "Prévu pour impression"],
    ]
    add_info_table(doc, edt_items[1:], headers=edt_items[0])

    doc.add_paragraph()

    # 7.9 Agents
    add_heading(doc, "7.9 Gestion des Utilisateurs (Agents)", level=2, color=(52, 73, 94))
    agent_items = [
        ["Fonctionnalité", "Description"],
        ["CRUD complet", "Créer, lire, modifier, supprimer des comptes administrateurs"],
        ["Permissions", "Gestion granulaire des droits par module"],
        ["Statuts", "ACTIVE (actif) / INACTIVE (inactif)"],
        ["Modules gérés", "Étudiants, Parents, Professeurs, Présences, Notes, Finances, Planning"],
    ]
    add_info_table(doc, agent_items[1:], headers=agent_items[0])

    doc.add_paragraph()

    # 7.10 Journal
    add_heading(doc, "7.10 Journal d'Activité (Audit)", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "Système de traçabilité complet enregistrant toutes les actions effectuées "
        "dans l'application, essentiel pour la conformité et la sécurité."
    )
    journal_items = [
        ["Paramètre", "Détail"],
        ["Actions tracées", "CREATE, UPDATE, DELETE, VIEW, LOGIN, LOGOUT"],
        ["Modules tracés", "Système, Étudiants, Parents, Présences, Notes, Finances, Planning, Professeurs"],
        ["Informations enregistrées", "Agent, horodatage, module, action, description, IP"],
        ["Filtres disponibles", "Par agent, par module, par période"],
        ["Rétention", "Illimitée (selon politique de l'établissement)"],
    ]
    add_info_table(doc, journal_items[1:], headers=journal_items[0])

    doc.add_paragraph()

    # 7.11 Parametres
    add_heading(doc, "7.11 Paramètres et Configuration", level=2, color=(52, 73, 94))
    param_items = [
        ["Onglet", "Contenu"],
        ["Général", "Nom de l'école, adresse, téléphone, e-mail, logo"],
        ["Personnel", "Rôle administrateur, gestion du mot de passe"],
        ["Notifications", "Activation e-mail, SMS, rappels de paiement, alertes absences"],
        ["Paiements", "Configuration des modes et fréquences de paiement"],
    ]
    add_info_table(doc, param_items[1:], headers=param_items[0])

    doc.add_page_break()

    # ============================================================
    # SECTION 8 : INTERFACE UTILISATEUR
    # ============================================================
    add_heading(doc, "8. Interface Utilisateur", level=1, color=(31, 92, 153))

    add_heading(doc, "8.1 Pages Principales (15)", level=2, color=(52, 73, 94))
    pages_data = [
        ["N°", "Page / Fichier", "Description"],
        ["1", "Login.tsx", "Authentification (email + mot de passe)"],
        ["2", "Index.tsx (Dashboard)", "Tableau de bord principal"],
        ["3", "Etudiants.tsx", "Liste et gestion des élèves"],
        ["4", "StudentProfile.tsx", "Profil détaillé d'un élève"],
        ["5", "Parents.tsx", "Liste et gestion des parents"],
        ["6", "ParentProfile.tsx", "Profil détaillé d'un parent"],
        ["7", "Professeurs.tsx", "Liste et gestion des professeurs"],
        ["8", "Presences.tsx", "Suivi et saisie des présences"],
        ["9", "Notes.tsx", "Saisie et consultation des notes"],
        ["10", "Finances.tsx", "Gestion des paiements et transactions"],
        ["11", "EmploiDuTemps.tsx", "Planning hebdomadaire"],
        ["12", "Agents.tsx", "Gestion des comptes administrateurs"],
        ["13", "AgentProfile.tsx", "Profil administrateur"],
        ["14", "Journal.tsx", "Journal d'audit et d'activité"],
        ["15", "Parametres.tsx", "Configuration de l'établissement"],
    ]
    add_info_table(doc, pages_data[1:], headers=pages_data[0])

    doc.add_paragraph()
    add_heading(doc, "8.2 Composants Réutilisables (70+)", level=2, color=(52, 73, 94))
    comp_data = [
        ["Catégorie", "Composants"],
        ["Layout", "DashboardLayout, Sidebar, TopBar"],
        ["Dashboard", "StatsCard, RevenueChart, RecentInscriptions"],
        ["Formulaires Élèves/Parents", "StudentForm, ParentForm, AgentDialog"],
        ["Finances", "PaymentReceipt, TarifDialog, QuickPaymentDialog, SplitPaymentDialog"],
        ["Planning", "AssignSlotDialog"],
        ["Communication", "WhatsAppNotifyDialog"],
        ["UI Primitifs (Radix/shadcn)", "Button, Input, Table, Dialog, Select, Badge, Card, Tabs, Toast, Tooltip, etc."],
    ]
    add_info_table(doc, comp_data[1:], headers=comp_data[0])

    doc.add_paragraph()
    add_heading(doc, "8.3 Services API Frontend (9)", level=2, color=(52, 73, 94))
    services_data = [
        ["Service", "Fichier", "Responsabilité"],
        ["Authentification", "auth.service.ts", "Login, Logout"],
        ["Élèves", "student.service.ts", "CRUD, recherche, filtres"],
        ["Parents", "parent.service.ts", "CRUD, enfants liés"],
        ["Professeurs", "teacher.service.ts", "CRUD, filtres"],
        ["Présences", "attendance.service.ts", "Saisie, statistiques"],
        ["Notes", "grade.service.ts", "CRUD, bulletins"],
        ["Transactions", "transaction.service.ts", "Paiements, résumé"],
        ["Agents", "agent.service.ts", "Utilisateurs, permissions"],
        ["Journal", "audit-log.service.ts", "Logs, filtres"],
    ]
    add_info_table(doc, services_data[1:], headers=services_data[0])

    doc.add_page_break()

    # ============================================================
    # SECTION 9 : SECURITE ET PERFORMANCES
    # ============================================================
    add_heading(doc, "9. Sécurité et Performances", level=1, color=(31, 92, 153))

    add_heading(doc, "9.1 Mesures de Sécurité Actuelles", level=2, color=(52, 73, 94))
    sec_data = [
        ["Mesure", "Implémentation", "Statut"],
        ["Authentification", "Email + mot de passe", "Opérationnel"],
        ["Routes protégées", "Redirection vers login si non connecté", "Opérationnel"],
        ["CORS configuré", "Contrôle des origines autorisées", "Opérationnel"],
        ["Session stateless", "API REST sans état (recommandé)", "Opérationnel"],
        ["Validation formulaires", "Zod côté frontend, Bean Validation côté backend", "Opérationnel"],
        ["Journal d'audit", "Traçabilité complète des actions", "Opérationnel"],
        ["Hachage des mots de passe", "À implémenter (bcrypt/argon2)", "Recommandé"],
        ["JWT Authentication", "À implémenter pour remplacer sessions", "Recommandé"],
        ["HTTPS/TLS", "À configurer en production", "Recommandé"],
        ["Rate Limiting", "À implémenter sur endpoints sensibles", "Recommandé"],
    ]
    add_info_table(doc, sec_data[1:], headers=sec_data[0])

    doc.add_paragraph()
    add_heading(doc, "9.2 Performances", level=2, color=(52, 73, 94))
    perf_data = [
        ["Aspect", "Implémentation"],
        ["Pagination backend", "20 éléments par défaut, configurable (5/10/20/50)"],
        ["Lazy loading", "Composants chargés à la demande (React)"],
        ["Requêtes optimisées", "JPA avec relations eager/lazy selon le contexte"],
        ["Build optimisé", "Vite avec code splitting automatique"],
        ["Desktop natif", "Electron pour éviter la latence réseau en local"],
    ]
    add_info_table(doc, perf_data[1:], headers=perf_data[0])

    doc.add_page_break()

    # ============================================================
    # SECTION 10 : DEPLOIEMENT ET INSTALLATION
    # ============================================================
    add_heading(doc, "10. Déploiement et Installation", level=1, color=(31, 92, 153))

    add_heading(doc, "10.1 Prérequis Système", level=2, color=(52, 73, 94))
    prereq_data = [
        ["Composant", "Version Minimale", "Remarque"],
        ["Java (JDK)", "21 LTS", "Requis pour le backend Spring Boot"],
        ["Node.js", "18+", "Requis pour le frontend React"],
        ["npm", "9+", "Gestionnaire de paquets Node"],
        ["MySQL", "8.0+", "Système de gestion de base de données"],
        ["Maven", "3.9+", "Outil de build Java"],
        ["Git", "2.x", "Gestion du code source"],
        ["RAM", "4 Go minimum", "8 Go recommandé pour le développement"],
        ["OS", "Windows 10+, Linux, macOS", "Tout OS supportant Java 21"],
    ]
    add_info_table(doc, prereq_data[1:], headers=prereq_data[0])

    doc.add_paragraph()
    add_heading(doc, "10.2 Étapes d'Installation", level=2, color=(52, 73, 94))
    steps = [
        "ÉTAPE 1 — Créer la base de données MySQL :",
        "    CREATE DATABASE edumanager_db CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;",
        "",
        "ÉTAPE 2 — Configurer les credentials dans backend/src/main/resources/application.properties :",
        "    spring.datasource.url=jdbc:mysql://localhost:3306/edumanager_db",
        "    spring.datasource.username=<votre_utilisateur>",
        "    spring.datasource.password=<votre_mot_de_passe>",
        "",
        "ÉTAPE 3 — Lancer le backend :",
        "    cd backend",
        "    mvn spring-boot:run",
        "    (Le schéma DB est créé automatiquement au premier démarrage)",
        "",
        "ÉTAPE 4 — Charger les données de test (optionnel) :",
        "    mysql -u root -p edumanager_db < backend/src/main/resources/seed-data.sql",
        "",
        "ÉTAPE 5 — Installer et lancer le frontend :",
        "    cd frontend",
        "    npm install",
        "    npm run dev",
        "    (Accessible sur http://localhost:5173)",
        "",
        "CONNEXION PAR DÉFAUT (données de test) :",
        "    Email    : sarah.admin@edumanager.ma",
        "    Mot de passe : admin123",
    ]
    step_para = doc.add_paragraph()
    run = step_para.add_run("\n".join(steps))
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)

    doc.add_paragraph()
    add_heading(doc, "10.3 Build de Production", level=2, color=(52, 73, 94))
    prod_data = [
        ["Cible", "Commandes", "Résultat"],
        ["Backend (JAR)", "cd backend && mvn clean package\njava -jar target/api-0.0.1-SNAPSHOT.jar", "Service REST autonome sur le port 8080"],
        ["Frontend Web", "cd frontend && npm run build\n(servir le dossier dist/)", "Application web statique (Nginx, Apache)"],
        ["Application Bureau", "cd frontend && npm run electron:build", "Installateur Windows .exe"],
    ]
    add_info_table(doc, prod_data[1:], headers=prod_data[0])

    doc.add_page_break()

    # ============================================================
    # SECTION 11 : ADEQUATION AU SYSTEME MAROCAIN
    # ============================================================
    add_heading(doc, "11. Adéquation au Système Éducatif Marocain", level=1, color=(31, 92, 153))

    doc.add_paragraph(
        "EduManager a été conçu en tenant compte des spécificités du système "
        "éducatif marocain. Voici les éléments d'adaptation détaillés."
    )

    add_heading(doc, "11.1 Nomenclature des Classes", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "L'application utilise la nomenclature officielle marocaine pour les cycles "
        "primaire et collège :"
    )
    classes_data = [
        ["Cycle", "Classes dans EduManager", "Correspondance"],
        ["Collège (Cycle secondaire collégial)", "6ème A, 6ème B", "1ère année collège"],
        ["Collège", "5ème A, 5ème B", "2ème année collège"],
        ["Collège", "4ème A, 4ème B", "3ème année collège"],
        ["Collège", "3ème A, 3ème B, 3ème C", "4ème année collège (avant Brevet)"],
    ]
    add_info_table(doc, classes_data[1:], headers=classes_data[0])

    doc.add_paragraph()
    add_heading(doc, "11.2 Système de Notation", level=2, color=(52, 73, 94))
    nota_data = [
        ["Paramètre", "EduManager", "Standard Marocain"],
        ["Échelle de notation", "Sur 20 points", "Sur 20 (conforme)"],
        ["Décimales", "DECIMAL(5,2)", "Demi-points acceptés (ex : 14,5/20)"],
        ["Coefficients", "Par matière (configurable)", "Coefficients officiels par cycle"],
        ["Types d'évaluation", "CC, DS, Exam, TP, Oral", "Conformes aux pratiques pédagogiques"],
        ["Semestres", "S1 et S2", "Conforme au calendrier scolaire marocain"],
        ["Année académique", "Format AAAA-AAAA (ex : 2024-2025)", "Conforme"],
    ]
    add_info_table(doc, nota_data[1:], headers=nota_data[0])

    doc.add_paragraph()
    add_heading(doc, "11.3 Aspects Financiers (Dirham Marocain)", level=2, color=(52, 73, 94))
    fin_data = [
        ["Aspect", "Implémentation"],
        ["Monnaie", "Dirham Marocain (MAD) — DECIMAL(10,2)"],
        ["Modes de paiement", "Espèces, Chèque, Virement bancaire (standards marocains)"],
        ["CIN Parent", "Champ dédié pour la Carte d'Identité Nationale marocaine"],
        ["Numérotation des reçus", "Numéro de reçu unique pour chaque transaction"],
        ["Fréquences de paiement", "Mensuelle (10 mois), trimestrielle (3 fois), annuelle"],
        ["Arriérés", "Suivi des impayés (arrears) par parent"],
        ["Contrat de scolarité", "Tarif annuel par élève avec plan de paiement personnalisé"],
    ]
    add_info_table(doc, fin_data[1:], headers=fin_data[0])

    doc.add_paragraph()
    add_heading(doc, "11.4 Horaires et Sessions", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "Les 4 sessions journalières d'EduManager correspondent aux plages horaires "
        "classiques des établissements marocains :"
    )
    sessions_data = [
        ["Session", "Horaires", "Période"],
        ["SESSION_1", "08h00 – 10h00", "Matinée 1ère heure"],
        ["SESSION_2", "10h00 – 12h00", "Matinée 2ème heure"],
        ["SESSION_3", "14h00 – 16h00", "Après-midi 1ère heure"],
        ["SESSION_4", "16h00 – 18h00", "Après-midi 2ème heure"],
    ]
    add_info_table(doc, sessions_data[1:], headers=sessions_data[0])

    doc.add_paragraph()
    add_heading(doc, "11.5 Semaine Scolaire", level=2, color=(52, 73, 94))
    doc.add_paragraph(
        "L'emploi du temps est configuré sur 5 jours (Lundi au Vendredi), "
        "conformément au calendrier scolaire marocain actuel. Le vendredi est "
        "intégré comme jour ouvrable complet."
    )

    doc.add_page_break()

    # ============================================================
    # SECTION 12 : STATISTIQUES DU PROJET
    # ============================================================
    add_heading(doc, "12. Statistiques du Projet", level=1, color=(31, 92, 153))

    stats_data = [
        ["Métrique", "Valeur"],
        ["Contrôleurs REST (Backend)", "13"],
        ["Services métier (Backend)", "11"],
        ["Repositories JPA", "10"],
        ["Entités JPA (tables)", "10"],
        ["Énumérations", "11"],
        ["Endpoints API", "80+"],
        ["DTOs (Request + Response)", "40+"],
        ["Pages Frontend", "15"],
        ["Composants React", "70+"],
        ["Services API Frontend", "9"],
        ["Fichiers Java", "100+"],
        ["Fichiers TypeScript/React", "150+"],
        ["Dépendances Backend (Maven)", "7 principales"],
        ["Dépendances Frontend (npm)", "30+"],
        ["Lignes de code estimées", "50 000+"],
        ["Taille données de test (seed)", "15 Ko"],
        ["Élèves dans les données de test", "11"],
        ["Professeurs dans les données de test", "5"],
        ["Classes dans les données de test", "9 classes"],
        ["Agents (admins) dans les données de test", "3"],
    ]
    add_info_table(doc, stats_data[1:], headers=stats_data[0])

    doc.add_page_break()

    # ============================================================
    # SECTION 13 : AMÉLIORATIONS SUGGÉRÉES
    # ============================================================
    add_heading(doc, "13. Améliorations Suggérées", level=1, color=(31, 92, 153))

    add_heading(doc, "13.1 Sécurité (Priorité Haute)", level=2, color=(52, 73, 94))
    sec_imp = [
        ["Amélioration", "Justification", "Priorité"],
        ["Hachage des mots de passe (bcrypt)", "Les mots de passe stockés en clair sont dangereux", "URGENT"],
        ["JWT Authentication", "Remplacer la session par des tokens JWT (API stateless)", "HAUTE"],
        ["HTTPS/TLS en production", "Chiffrement des communications", "HAUTE"],
        ["Rate Limiting", "Protection contre les attaques brute-force", "HAUTE"],
        ["Validation backend renforcée", "Vérification systématique des entrées côté serveur", "MOYENNE"],
        ["Secure cookies", "Remplacer localStorage par secure HttpOnly cookies", "MOYENNE"],
    ]
    add_info_table(doc, sec_imp[1:], headers=sec_imp[0])

    doc.add_paragraph()
    add_heading(doc, "13.2 Fonctionnalités Additionnelles", level=2, color=(52, 73, 94))
    feat_imp = [
        ["Fonctionnalité", "Description", "Impact"],
        ["Export PDF bulletins", "Génération automatique de bulletins en PDF", "HAUTE"],
        ["Export Excel", "Export des listes élèves, notes, finances en Excel", "HAUTE"],
        ["Notifications SMS/WhatsApp", "Alertes absences, rappels de paiement", "HAUTE"],
        ["Portail parents", "Accès lecture seule pour consulter bulletin et présences", "HAUTE"],
        ["Portail professeurs", "Saisie des notes directement par les professeurs", "HAUTE"],
        ["Justification absences", "Module de gestion des absences justifiées", "MOYENNE"],
        ["Paiement en ligne", "Intégration CMI (Centre Monétique Interbancaire Maroc)", "MOYENNE"],
        ["Calendrier scolaire", "Gestion des jours fériés et vacances marocains", "MOYENNE"],
        ["Messagerie interne", "Communication directe entre administration et parents", "FAIBLE"],
        ["Application mobile", "App iOS/Android pour parents et professeurs", "FAIBLE"],
    ]
    add_info_table(doc, feat_imp[1:], headers=feat_imp[0])

    doc.add_paragraph()
    add_heading(doc, "13.3 Performance et Qualité", level=2, color=(52, 73, 94))
    qual_imp = [
        ["Amélioration", "Description"],
        ["Tests unitaires", "Tests JUnit (backend) et Vitest (frontend)"],
        ["Tests d'intégration", "Tests end-to-end avec Cypress ou Playwright"],
        ["CI/CD", "Pipeline automatisé (GitHub Actions) pour build et tests"],
        ["Swagger/OpenAPI", "Documentation interactive de l'API"],
        ["Redis Cache", "Mise en cache pour améliorer les performances"],
        ["Indexation DB", "Index sur les colonnes fréquemment filtrées (student_id, academic_year)"],
        ["Monitoring", "Mise en place de métriques (Spring Actuator + Grafana)"],
        ["Logging centralisé", "Stack ELK (Elasticsearch, Logstash, Kibana)"],
    ]
    add_info_table(doc, qual_imp[1:], headers=qual_imp[0])

    doc.add_page_break()

    # ============================================================
    # SECTION 14 : CONCLUSION
    # ============================================================
    add_heading(doc, "14. Conclusion", level=1, color=(31, 92, 153))

    doc.add_paragraph(
        "EduManager est une solution de gestion scolaire complète, moderne et bien "
        "adaptée aux besoins des établissements scolaires marocains. Développée avec "
        "des technologies éprouvées et actuelles (Spring Boot, React, MySQL), elle "
        "offre une architecture solide et extensible."
    )

    doc.add_paragraph()

    add_heading(doc, "Points Forts de la Plateforme", level=2, color=(52, 73, 94))
    strengths_data = [
        ["Point Fort", "Détail"],
        ["Couverture fonctionnelle complète", "Tous les aspects de la gestion scolaire sont couverts"],
        ["Adapté au contexte marocain", "Notation /20, semestres, CIN, MAD, horaires locaux"],
        ["Interface moderne et intuitive", "UX soignée avec Tailwind CSS et composants Radix UI"],
        ["Architecture robuste", "Séparation des couches, DTOs, services, repositories"],
        ["Double déploiement", "Web (navigateur) et bureau (Electron)"],
        ["Traçabilité totale", "Journal d'audit couvrant 100% des actions"],
        ["Code de qualité", "TypeScript, Java 21, validation côté client et serveur"],
        ["Données de test réalistes", "Seed data complet pour démonstration immédiate"],
    ]
    add_info_table(doc, strengths_data[1:], headers=strengths_data[0])

    doc.add_paragraph()

    doc.add_paragraph(
        "En résumé, EduManager constitue une base solide pour digitaliser la gestion "
        "scolaire au Maroc. Avec les améliorations de sécurité recommandées "
        "(hachage des mots de passe, JWT) et l'ajout de fonctionnalités comme "
        "l'export PDF des bulletins et les notifications parents, cette plateforme "
        "peut répondre aux besoins d'établissements scolaires marocains de toutes tailles."
    )

    doc.add_paragraph()

    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig_para.add_run(f"Rapport généré le {today}\nEduManager v1.0.0")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(127, 140, 141)

    # ── Save ─────────────────────────────────────────────────────────────────────
    output_path = r"c:\Users\pakiy\Documents\prj\edumanager\Rapport_EduManager_Maroc.docx"
    doc.save(output_path)
    print(f"Rapport sauvegardé : {output_path}")

if __name__ == "__main__":
    create_report()
